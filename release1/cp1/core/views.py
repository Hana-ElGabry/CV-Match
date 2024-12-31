from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from .factories.document_creator import DocumentFactory
from io import BytesIO  # Required for BytesIO in PDF generation
from reportlab.pdfgen import canvas  # for PDF creation
import docx  # for Word document creation

# Home view
def Home(request):
    return render(request, 'index.html')

# Register view
def RegisterView(request):
    if request.method == "POST":
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        username = request.POST.get('username')
        email = request.POST.get('email')
        password = request.POST.get('password')

        user_data_has_error = False

        if User.objects.filter(username=username).exists():
            user_data_has_error = True
            messages.error(request, "Username already exists")

        if User.objects.filter(email=email).exists():
            user_data_has_error = True
            messages.error(request, "Email already exists")

        if len(password) < 5:
            user_data_has_error = True
            messages.error(request, "Password must be at least 5 characters")

        if user_data_has_error:
            return redirect('register')
        else:
            new_user = User.objects.create_user(
                first_name=first_name,
                last_name=last_name,
                email=email, 
                username=username,
                password=password
            )
            messages.success(request, "Account created. Login now")
            return redirect('login')

    return render(request, 'register.html')

# Login view
def LoginView(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            return redirect('profile')
        
        else:
            messages.error(request, "Invalid login credentials")
            return redirect('login')

    return render(request, 'login.html')

# Profile view (requires login)
@login_required
def profile(request):
    return render(request, 'profile.html')

# Document download views

def download_document(request, doc_type):
    factory = DocumentFactory()
    document = factory.create_document(doc_type)

    if not document:
        return HttpResponse("Invalid document type.", status=400)

    response = HttpResponse(content_type=document.content_type)
    response['Content-Disposition'] = f'attachment; filename="profile.{doc_type}"'
    response.write(document.generate(request.user))  # Assuming the generate method creates the document content
    return response

def download_pdf(request, user_id):
    user = User.objects.get(id=user_id)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="user_profile.pdf"'

    buffer = BytesIO()  # BytesIO is used to create in-memory PDF
    p = canvas.Canvas(buffer)
    p.drawString(100, 750, f"Username: {user.username}")
    p.drawString(100, 730, f"Email: {user.email}")
    p.drawString(100, 710, f"Joined: {user.date_joined}")
    p.showPage()
    p.save()

    response.write(buffer.getvalue())
    return response

def download_word(request, user_id):
    user = User.objects.get(id=user_id)
    doc = docx.Document()
    doc.add_paragraph(f"Username: {user.username}")
    doc.add_paragraph(f"Email: {user.email}")
    doc.add_paragraph(f"Joined: {user.date_joined}")

    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename="user_profile.docx"'
    doc.save(response)
    return response

