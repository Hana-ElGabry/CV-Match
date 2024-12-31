from abc import ABC, abstractmethod
from xhtml2pdf import pisa
from django.template.loader import render_to_string
from io import BytesIO
from docx import Document


class DocumentCreator(ABC):
    @abstractmethod
    def create_document(self, context):
        pass


class PDFDocumentCreator(DocumentCreator):
    def create_document(self, context):
        template_name = context.get('template_name')
        html_content = render_to_string(template_name, context)

        result = BytesIO()
        pisa_status = pisa.CreatePDF(html_content, dest=result)

        if pisa_status.err:
            raise Exception("Error generating PDF document")

        result.seek(0)
        return result


class WordDocumentCreator(DocumentCreator):
    def create_document(self, context):
        # Create a new Word document
        document = Document()

        # Add content from the context
        document.add_heading(f"{context['user'].username}'s Profile", level=1)

        if context['user'].first_name:
            document.add_paragraph(f"First Name: {context['user'].first_name}")
        if context['user'].last_name:
            document.add_paragraph(f"Last Name: {context['user'].last_name}")
        if context['user'].email:
            document.add_paragraph(f"Email: {context['user'].email}")
        if context.get('profile', {}).get('job_description'):
            document.add_paragraph(f"Job Description: {context['profile']['job_description']}")

        result = BytesIO()
        document.save(result)
        result.seek(0)
        return result


class DocumentFactory:
    @staticmethod
    def get_creator(doc_type):
        if doc_type == 'pdf':
            return PDFDocumentCreator()
        elif doc_type == 'word':
            return WordDocumentCreator()
        else:
            raise ValueError("Invalid document type")
